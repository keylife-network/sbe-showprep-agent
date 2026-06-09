#!/usr/bin/env python3
"""
SBE Show Prep Generator
Creates weekly show prep documents for Steve Brown, Etc. on Key Life Network.

Usage:
    python3 generate_show_prep.py --author "Author Name" --book "Book Title" [options]

Options:
    --episode NUM       Episode number (e.g. 28)
    --airdate DATE      Air date as M.DD.YY (e.g. 7.12.26)
    --year YEAR         Year directory (default: current year)
    --no-search         Skip web search (faster, less accurate socials)
    --output-dir DIR    Output directory override
    --api-key KEY       Anthropic API key (or set ANTHROPIC_API_KEY env var, or put it in .env)

SETUP:
    Get an API key from https://console.anthropic.com/
    Then either:
      export ANTHROPIC_API_KEY=sk-ant-...   (add to ~/.zshrc)
      or create a .env file in this directory with: ANTHROPIC_API_KEY=sk-ant-...
"""

import argparse
import json
import os
import re
import shutil
import sys
from datetime import datetime, timedelta
from pathlib import Path

# Load .env file if present (before importing anthropic)
_env_file = Path(__file__).parent / ".env"
if _env_file.exists():
    for line in _env_file.read_text().splitlines():
        line = line.strip()
        if line and not line.startswith("#") and "=" in line:
            k, _, v = line.partition("=")
            os.environ.setdefault(k.strip(), v.strip())

import anthropic
from docx import Document
from docx.shared import Pt, Inches
from duckduckgo_search import DDGS


# ── Web Search ────────────────────────────────────────────────────────────────

def search_web(query: str, max_results: int = 4) -> list[dict]:
    try:
        with DDGS() as ddgs:
            return list(ddgs.text(query, max_results=max_results))
    except Exception as e:
        print(f"  [search] {e}", file=sys.stderr)
        return []


def gather_author_info(author: str, book: str) -> str:
    print(f"  Searching web for: {author} / {book} ...")
    searches = [
        f"{author} author biography Christian",
        f"{author} {book}",
        f"{author} Twitter X Instagram Facebook social media",
        f"{author} website ministry podcast church",
    ]
    parts = []
    for q in searches:
        results = search_web(q, max_results=3)
        if results:
            parts.append(f"\n--- Search: '{q}' ---")
            for r in results:
                parts.append(f"Title: {r.get('title', '')}")
                parts.append(f"URL:   {r.get('href', '')}")
                body = r.get("body", "")[:400]
                if body:
                    parts.append(f"Body:  {body}")
                parts.append("")
    return "\n".join(parts) if parts else "(no results)"


# ── Claude Content Generation ─────────────────────────────────────────────────

SHOW_CONTEXT = """
SHOW: Steve Brown, Etc. on Key Life Network
HOST: Dr. Steve Brown — 80-something, witty Reformed theologian, been in radio 40+ years.
Famous for grace-centered preaching, self-deprecating humor, and talking to real broken people.

REGULAR CAST:
- Matthew Porter: Steve's co-host
- Jeremy: Producer in "the little glass booth"
- John Myers: One-man IT department in "the tech bunker"
- Dr. George Bingham: President of Key Life
- Cathy Wyatt: "The soft, feminine side of the program"

TONE: Warm, theologically serious but accessible, genuinely funny.
The audience is Christians who want theological depth AND real-life application.
Steve always circles back to grace, freedom, and the gospel.
"""


def generate_content(author: str, book: str, web_context: str) -> dict:
    client = anthropic.Anthropic()
    print("  Generating content with Claude ...")

    prompt = f"""{SHOW_CONTEXT}

GUEST:  {author}
BOOK/PROJECT: {book}

WEB RESEARCH RESULTS:
{web_context}

Generate show prep. Return ONLY a single valid JSON object — no markdown fences, no extra text.
Use the author's real information from the web research where possible.

JSON keys required:

{{
  "opening_teaser": "ALL CAPS 1-2 sentence dramatic hook about the topic. End with 'LET’S TALK ABOUT IT WITH {author.upper()}… ON STEVE BROWN, ETC.'",

  "matthew_banter": "Full line. 'Matthew Porter is here. [genuinely funny joke, often riffing on the book topic]. [Matthew responds]'",
  "jeremy_banter": "Full line. 'Our producer, Jeremy, is in the little glass booth. [witty one-liner]'",
  "john_banter": "Full line. 'Our one-man IT department, John Myers, is in the tech bunker. [tech/IT humor]'",
  "george_banter": "Full line. 'Dr. George Bingham is the president of Key Life. [quip — often coffee or ministry]'",
  "cathy_banter": "Full line. 'And Cathy Wyatt is the soft, feminine side of the program. [gentle wry banter]'",

  "guest_bio": "2-3 sentence radio intro. First word(s) should be the author's name. End with '[First/Last name]’s new book is called [Book Title].'",

  "segment_1_questions": [
    "Question 1 — background / origin story",
    "Question 2 — faith / spiritual journey",
    "Question 3 — core idea of the book",
    "Question 4 — key insight or surprising point",
    "Question 5 — provocative follow-up"
  ],

  "segment_2_rejoin": "One sentence. 'Thanks for joining us on Steve Brown, Etc. We’re talking to [descriptor] {author}. [His/Her] new book is called {book}.'",
  "segment_2_questions": [
    "Deeper dive into book themes",
    "Grace/gospel connection",
    "Something counterintuitive from the book",
    "Personal application"
  ],

  "segment_3_rejoin": "Rejoin + social mention. 'You’re listening to Steve Brown, Etc. and we’re hanging out with [descriptor] {author}. [Keep up with him/her at website and on social platforms with handles from research — or ask if unknown]'",
  "segment_3_questions": [
    "What’s the one takeaway you want listeners to walk away with?",
    "One final thematic question"
  ],

  "segment_4_questions": [
    "Personal faith or current spiritual state question",
    "A Steve Brown-style closing question — about grace, assurance, or hope for broken people"
  ],

  "social_handles": {{
    "x_twitter": ["@handle1", "@publisherHandle"],
    "facebook": ["@handle"],
    "instagram": ["@handle"],
    "youtube": ["@handle"],
    "tiktok": [],
    "website": "author-website.com"
  }},

  "hashtags": ["#AuthorLastName", "#BookTitleNoSpaces", "#TopicTag", "#PublisherIfKnown"],

  "yt_title": "Compelling hook | {author} | Steve Brown, Etc.",
  "yt_description_hook": "2-3 sentence YouTube opener. Engaging hook about the topic. Mention book title and guest.",
  "yt_guest_bio": "Paragraph-length YouTube guest bio with credentials, ministry, and social @ handles.",
  "yt_hashtags": "#Tag1 #Tag2 #Tag3 #Tag4 #Tag5"
}}"""

    message = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=5000,
        messages=[{"role": "user", "content": prompt}],
    )

    raw = message.content[0].text.strip()

    # Strip markdown fences if present
    raw = re.sub(r"^```(?:json)?\s*", "", raw)
    raw = re.sub(r"\s*```$", "", raw)

    # Extract outermost JSON object
    m = re.search(r"(\{.*\})", raw, re.DOTALL)
    if not m:
        print("ERROR: No JSON found in Claude response", file=sys.stderr)
        print(raw[:800], file=sys.stderr)
        sys.exit(1)

    try:
        return json.loads(m.group(1))
    except json.JSONDecodeError as e:
        print(f"ERROR: JSON parse failed — {e}", file=sys.stderr)
        print(m.group(1)[:800], file=sys.stderr)
        sys.exit(1)


# ── Document Builder ──────────────────────────────────────────────────────────

PT18 = Pt(18)


def _set_size(run):
    run.font.size = PT18


def add_blank(doc):
    p = doc.add_paragraph()
    return p


def add_normal(doc, text, size=True):
    p = doc.add_paragraph()
    r = p.add_run(text)
    if size:
        _set_size(r)
    return p


def add_bold(doc, text, size=True):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    if size:
        _set_size(r)
    return p


def add_bio(doc, author, bio_text, size=True):
    """Bio paragraph: author name bold, rest normal."""
    p = doc.add_paragraph()
    r1 = p.add_run(author)
    r1.bold = True
    if size:
        _set_size(r1)
    rest = bio_text[len(author):] if bio_text.startswith(author) else " " + bio_text
    r2 = p.add_run(rest)
    if size:
        _set_size(r2)
    return p


def add_social_line(doc, platform, handles: list[str], size=True):
    """Platform label bold + first handle on same line; extra handles on new lines."""
    if not handles:
        return
    # Filter empty/placeholder handles
    handles = [h for h in handles if h and h not in ("@handle1", "@handle", "@publisherHandle", "@handle2")]
    if not handles:
        handles = ["@"]

    p = doc.add_paragraph()
    label_run = p.add_run(f"{platform} |\t\t\t\t")
    label_run.bold = True
    if size:
        _set_size(label_run)
    h_run = p.add_run(handles[0] if handles[0].startswith("@") else f"@{handles[0]}")
    if size:
        _set_size(h_run)

    for extra in handles[1:]:
        ep = doc.add_paragraph()
        er = ep.add_run(f"\t\t\t\t{extra if extra.startswith('@') else '@' + extra}")
        if size:
            _set_size(er)


def build_document(
    author: str,
    book: str,
    episode_num: str,
    airdate: str,
    content: dict,
    output_path: str,
):
    doc = Document()

    # Page layout: match template (18pt body, standard margins)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    socials = content.get("social_handles", {})

    # ── Opening Teaser ───────────────────────────────────────────────────────
    add_normal(doc, content["opening_teaser"])
    add_blank(doc)

    # ── Studio Intros ────────────────────────────────────────────────────────
    add_normal(doc, "Welcome to Steve Brown, Etc. I'm Steve, the aforementioned old white guy.")
    add_blank(doc)
    add_normal(doc, content["matthew_banter"])
    add_blank(doc)
    add_normal(doc, content["jeremy_banter"])
    add_blank(doc)
    add_normal(doc, content["john_banter"])
    add_blank(doc)
    add_normal(doc, content["george_banter"])
    add_blank(doc)
    add_normal(doc, content["cathy_banter"])

    # ── Guest Intro ──────────────────────────────────────────────────────────
    add_bold(doc, f"Intro Guest - {author}")
    add_blank(doc)
    bio = content["guest_bio"]
    add_bio(doc, author, bio)
    add_blank(doc)

    # ── Segment 1 ────────────────────────────────────────────────────────────
    add_bold(doc, "SEGMENT 1 — 10:20")
    for q in content.get("segment_1_questions", []):
        add_normal(doc, q)
        add_blank(doc)
    add_blank(doc)

    # ── Segment 2 ────────────────────────────────────────────────────────────
    add_bold(doc, "SEGMENT 2 — 7:50")
    add_bold(doc, "Rejoin")
    add_normal(doc, content["segment_2_rejoin"])
    add_blank(doc)
    for q in content.get("segment_2_questions", []):
        add_normal(doc, q)
        add_blank(doc)
    add_blank(doc)

    # ── Segment 3 ────────────────────────────────────────────────────────────
    add_bold(doc, "SEGMENT 3 — 5:50")
    add_bold(doc, "Rejoin")
    add_normal(doc, content["segment_3_rejoin"])
    add_blank(doc)
    for q in content.get("segment_3_questions", []):
        add_normal(doc, q)
        add_blank(doc)
    add_blank(doc)

    # ── Segment 4 ────────────────────────────────────────────────────────────
    add_bold(doc, "SEGMENT 4 — 7:50")
    add_bold(doc, "Rejoin")
    for q in content.get("segment_4_questions", []):
        add_normal(doc, q)
        add_blank(doc)

    seg4_plug = (
        "Thanks for spending time with us here on Steve Brown, Etc. "
        "And by the way, if you could use a reminder that God’s not mad at His children "
        "— especially on a Monday morning — we have just the thing… "
        "It’s called Key Life Connection and it’s our free weekly email. "
        "Check it out at KeyLife.org / Subscribe."
    )
    add_normal(doc, seg4_plug)
    add_blank(doc)
    add_blank(doc)

    # ── Segment 5 ────────────────────────────────────────────────────────────
    add_bold(doc, "SEGMENT 5 — 2:40")
    seg5_text = (
        "Rejoin" + " " * 85 +
        "Thanks for spending an hour with us here on Steve Brown, Etc. "
        "And in case you didn’t know, Key Life has added a new podcast called Simply Sermons. "
        "You can find it on the Key Life app, on your favorite podcast platform, "
        "or at KeyLife.org / Simply Sermons."
    )
    p5 = doc.add_paragraph()
    r5a = p5.add_run("Rejoin")
    r5a.bold = True
    _set_size(r5a)
    r5b = p5.add_run(
        " " * 85 +
        "Thanks for spending an hour with us here on Steve Brown, Etc. "
        "And in case you didn’t know, Key Life has added a new podcast called Simply Sermons. "
        "You can find it on the Key Life app, on your favorite podcast platform, "
        "or at KeyLife.org / Simply Sermons."
    )
    _set_size(r5b)
    add_blank(doc)

    # ── Wrapup ───────────────────────────────────────────────────────────────
    add_normal(doc, "Wrapup")
    add_normal(doc, "Tee up next week’s guest…")
    add_blank(doc)

    # ── Metadata ─────────────────────────────────────────────────────────────
    add_bold(doc, "Show Title")
    add_normal(doc, f"{author} | {book} | Steve Brown, Etc.")
    add_blank(doc)

    add_bold(doc, "Airdate")
    add_normal(doc, f"Weekend of {airdate} (posted at KeyLife.org on {airdate})")
    add_blank(doc)
    add_blank(doc)

    # Social media
    add_social_line(doc, "X", socials.get("x_twitter", []))
    add_blank(doc)
    add_social_line(doc, "Facebook", socials.get("facebook", []))
    add_blank(doc)
    add_social_line(doc, "Insta", socials.get("instagram", []))
    add_blank(doc)
    add_social_line(doc, "YouTube", socials.get("youtube", []))
    if socials.get("tiktok"):
        add_blank(doc)
        add_social_line(doc, "TikTok", socials.get("tiktok", []))
    add_blank(doc)
    add_blank(doc)

    # Hashtags
    add_bold(doc, "Hashtags")
    for tag in content.get("hashtags", []):
        add_normal(doc, tag)
    add_blank(doc)

    add_bold(doc, "YT Shorts Target Clip")
    add_blank(doc)
    add_blank(doc)

    add_bold(doc, "Etc. Insider")
    add_blank(doc)
    add_blank(doc)

    add_bold(doc, "Amazon Book")
    add_blank(doc)
    add_blank(doc)
    add_blank(doc)
    add_blank(doc)

    # ── YouTube Copy / Settings / Tags ────────────────────────────────────────
    add_bold(doc, "YouTube Copy / Settings / Tags")
    add_blank(doc)

    add_normal(doc, "TITLE")
    add_blank(doc)
    add_normal(doc, content.get("yt_title", f"{author} | {book} | Steve Brown, Etc."))
    add_blank(doc)

    add_normal(doc, "DESCRIPTION")
    add_blank(doc)
    add_normal(doc, content.get("yt_description_hook", ""))
    add_blank(doc)
    add_normal(doc, (
        "\U0001f5dd️ What do YOU think about this topic? "
        "Did this episode encourage your faith? Share your thoughts in the Comments. \U0001f5dd️"
    ))
    add_blank(doc)
    add_normal(doc, (
        "While you’re here, please help Key Life by clicking the ‘Subscribe’ button, "
        "then click on the BELL (on mobile devices, also click ‘ALL’). "
        "That way you’ll be the first to know when the newest episode of "
        "‘Steve Brown, Etc.’ goes live!"
    ))
    add_blank(doc)
    add_normal(doc, (
        "Get the best of Key Life in your inbox with Key Life Connection, our free weekly email: "
        "https://www.KeyLife.org/subscribe – includes FREE access to our "
        "behind-the-scenes ‘SBE Insider’ playlist here on YouTube."
    ))
    add_normal(doc, "And be sure to download the Key Life app, now available for iPhone and Android: https://www.KeyLife.org/app")
    add_blank(doc)

    add_normal(doc, "GUEST BIO")
    add_blank(doc)
    add_normal(doc, content.get("yt_guest_bio", ""))
    add_blank(doc)

    add_normal(doc, "00:00 - Intro")
    add_normal(doc, f"02:30 - Meet {author}")
    add_normal(doc, "xx:xx - Create blurb that summarizes guest’s first answer.")
    add_blank(doc)
    add_normal(doc, "THEN, skip forward every ten minutes and create a new chapter. There won’t be that many.")
    add_blank(doc)
    add_normal(doc, "37:38 - Steve’s closing thoughts")
    add_normal(doc, "39:52 - Next week’s guest is...")
    add_blank(doc)

    add_normal(doc, "Follow us on Social!")
    add_normal(doc, "▶ X/TWITTER – /keylifenetwork")
    add_normal(doc, "▶ FACEBOOK – /keylifenetwork")
    add_normal(doc, "▶ INSTAGRAM – /keylifenetwork")
    add_blank(doc)
    add_normal(doc, (
        "Key Life participates in the Amazon Services LLC Associates Program, "
        "an affiliate advertising program designed to provide a means for sites to earn "
        "advertising fees by advertising and linking to Amazon.com. "
        "Some links may be affiliate links. We may get a small referral fee if you buy "
        "something or take an action after clicking one of these."
    ))
    add_blank(doc)
    add_normal(doc, content.get("yt_hashtags", ""))
    add_blank(doc)
    add_blank(doc)

    # YouTube Settings (boilerplate)
    for label, value in [
        ("PLAYLISTS", "Steve Brown, Etc."),
        ("AUDIENCE", '"No, it\'s not made for kids"'),
        ("PAID PROMOTION", "Unchecked"),
        ("ALTERED CONTENT", "No."),
        ("AUTOMATED CHAPTERS", "Unchecked."),
        ("FEATURED PLACES", "Allow automatic places (checked)\nAllow automatic concepts (checked)"),
        ("TAGS", (
            "Key Life,Key Life Network,Steve Brown,Author Steve Brown,Dr. Steve Brown,"
            "grace,radical grace,grace teaching,christian talk radio,Christianity,Christian"
        )),
    ]:
        add_blank(doc)
        add_normal(doc, label)
        add_blank(doc)
        for line in value.split("\n"):
            add_normal(doc, line)

    doc.save(output_path)
    print(f"  Saved: {output_path}")


# ── CLI ───────────────────────────────────────────────────────────────────────

def next_episode_num(year_dir: Path) -> str:
    """Find the next episode number based on existing files."""
    existing = list(year_dir.glob("[0-9][0-9] - *_Show Prep.docx"))
    if not existing:
        return "01"
    nums = []
    for f in existing:
        m = re.match(r"^(\d+)", f.name)
        if m:
            nums.append(int(m.group(1)))
    return f"{max(nums) + 1:02d}"


def main():
    parser = argparse.ArgumentParser(
        description="Generate SBE show prep document",
        epilog=(
            "TIP: If you're already inside Claude Code, just say "
            "'Generate show prep for [Author] / [Book]' — no API key needed."
        ),
    )
    parser.add_argument("--author", required=True, help="Guest/author name")
    parser.add_argument("--book", required=True, help="Book or project title")
    parser.add_argument("--episode", default=None, help="Episode number (e.g. 28)")
    parser.add_argument("--airdate", default=None, help="Air date as M.DD.YY (e.g. 7.12.26)")
    parser.add_argument("--year", default=None, type=int, help="Year directory (default: current year)")
    parser.add_argument("--output-dir", default=None, help="Override output directory")
    parser.add_argument("--no-search", action="store_true", help="Skip web search")
    parser.add_argument("--api-key", default=None, help="Anthropic API key (overrides ANTHROPIC_API_KEY env var)")
    args = parser.parse_args()

    if args.api_key:
        os.environ["ANTHROPIC_API_KEY"] = args.api_key

    if not os.environ.get("ANTHROPIC_API_KEY"):
        print("ERROR: No Anthropic API key found.")
        print()
        print("Options:")
        print("  1. Pass it directly:    --api-key sk-ant-...")
        print("  2. Set env var:         export ANTHROPIC_API_KEY=sk-ant-...")
        print("  3. Create .env file:    echo 'ANTHROPIC_API_KEY=sk-ant-...' > .env")
        print("  4. Use Claude Code directly (no key needed): just ask Claude to generate show prep")
        print()
        print("Get a key at: https://console.anthropic.com/")
        sys.exit(1)

    base_dir = Path(__file__).parent
    year = args.year or datetime.now().year
    year_dir = Path(args.output_dir) if args.output_dir else base_dir / str(year)

    if not year_dir.exists():
        print(f"Creating directory: {year_dir}")
        year_dir.mkdir(parents=True)

    episode_num = args.episode or next_episode_num(year_dir)
    if args.airdate:
        airdate = args.airdate
    else:
        # Default to next weekend (Saturday)
        today = datetime.now()
        days_until_sat = (5 - today.weekday()) % 7
        if days_until_sat == 0:
            days_until_sat = 7
        next_sat = today + timedelta(days=days_until_sat)
        airdate = next_sat.strftime("%-m.%-d.%y")

    print(f"\nSBE Show Prep Generator")
    print(f"  Author:  {args.author}")
    print(f"  Book:    {args.book}")
    print(f"  Episode: {episode_num}")
    print(f"  Airdate: {airdate}")
    print(f"  Output:  {year_dir}/\n")

    # Step 1: Web search
    if args.no_search:
        web_context = "(web search skipped)"
    else:
        web_context = gather_author_info(args.author, args.book)

    # Step 2: Generate content
    content = generate_content(args.author, args.book, web_context)

    # Step 3: Build document
    safe_author = re.sub(r'[<>:"/\\|?*]', "", args.author)
    filename = f"{episode_num} - {safe_author}_Show Prep.docx"
    output_path = str(year_dir / filename)

    build_document(
        author=args.author,
        book=args.book,
        episode_num=episode_num,
        airdate=airdate,
        content=content,
        output_path=output_path,
    )

    print(f"\nDone! Open: {output_path}")
    print("\nRemember to manually fill in:")
    print("  - YT Shorts Target Clip timestamp")
    print("  - Etc. Insider timestamp")
    print("  - Amazon Book link")
    print("  - Posted date (the day after airdate)")
    print("  - Verify/correct social handles")


if __name__ == "__main__":
    main()
